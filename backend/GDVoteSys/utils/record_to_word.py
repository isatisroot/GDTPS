from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from apps.models import Meeting, ShareholderInfo, OnSiteMeeting, GB, MotionBook, AccumulateMotion
from GDVoteSys.settings import BASE_DIR
class RecordToWord():
    """
    将统计的唱票结果写入到word文档中
    """
    def __init__(self):
        self.document = Document()
    def vote_result(self,year, name):
        m = Meeting.objects.filter(year=year, name=name)[0]
        totalShare= m.gb.gb
        AShareTotal= m.gb.ltag
        BShareTotal=m.gb.ltbg
        qs = m.onsitemeeting_set.all()  # 该会议下所有等级的股东信息
        qs1 = qs.filter(cx=True)    # 出席会议的股东
        cx_num = len(qs1)
        cx_gb = 0  # 出席会议股东股份数总和
        for q in qs1:
            cx_gb += q.gzA + q.gzB
        percent1 = cx_gb / totalShare


        qs2 = qs.filter(cx=True, xcorwl=True)
        xcorwl_num = len(qs2)
        xcorwl_gb = 0
        for q in qs2:
            xcorwl_gb += q.gzA + q.gzB
        percent2 = xcorwl_gb / totalShare
        xcorwl_gb = format(xcorwl_gb, ",")

        qs3= qs.filter(cx=True, xcorwl=False)
        wl_num = len(qs3)
        wl_gb = 0
        for q in qs3:
            wl_gb += q.gzA + q.gzB
        percent3 = wl_gb / totalShare
        wl_gb = format(wl_gb, ",")

        # 设置标题
        title1 = self.document.add_heading(level=2)
        title2 = self.document.add_heading(level=1)
        # 标题居中
        title1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # 添加标题内容
        title1_run = title1.add_run("佛山电器照明股份有限公司")
        title2_run = title2.add_run("{}议案投票表决统计结果".format(year+name))
        self.document.add_heading( level=1)
        self.document.add_paragraph("一、会议的出席情况")
        self.document.add_paragraph("1. 出席的总体情况：")
        # format(cx_gb, ",")转千分符
        content1 = """        股东（代理人）{}人，代表股份{}股，占公司有表决权总股份{:.2%}。其中,现场会议股东（代理人）{}人，代表股份{}股，占公司有表决权总股份{:.2%}。网络投标股东（代理人）{}人，代表股份{}股，占公司有表决权总股份{:.2%}。"""\
            .format(cx_num, format(cx_gb, ","), percent1, xcorwl_num, xcorwl_gb, percent2, wl_num, wl_gb, percent3)
        self.document.add_paragraph(content1)
        self.document.add_paragraph("2. A股股东出席情况：")
        qs4 = qs.filter(gzA__gt=0, cx=True) # 出席A股股东
        num4 = len(qs4)
        sum4 = 0 # 出席会议所有股东的A股总数
        for q in qs4:
            sum4 += q.gzA
        percent4 = sum4 / AShareTotal
        # sum4 = format(sum4, ",")

        qs5 = qs.filter(gzA__gt=0, xcorwl=True)
        num5 = len(qs5)
        sum5 = 0
        for q in qs5:
            sum5 += q.gzA
        percent5 = sum5 / AShareTotal
        # sum5 = format(sum5, ",")

        qs6 = qs.filter(gzA__gt=0, xcorwl=False)
        num6 = len(qs6)
        sum6 = 0
        for q in qs6:
            sum6 += q.gzA
        percent6= sum6 / AShareTotal
        # sum6 = format(sum6, ",")

        content2 = """        A股股东（代理人）{}人，代表股份{}股，占公司A股股东表决权股份{:.2%}。其中,现场会议股东（代理人）{}人，代表股份{}股，占公司有表决权总股份{:.2%}。网络投标股东（代理人）{}人，代表股份{}股，占公司有表决权总股份{:.2%}。"""\
            .format(num4, sum4, percent4, num5, sum5, percent5, num6, sum6, percent6)
        self.document.add_paragraph(content2)

        self.document.add_paragraph("3. B股股东出席情况：")
        qs7 = qs.filter(gzB__gt=0, cx=True) # 出席B股股东
        num7 = len(qs7)
        num8 = num9 = sum7 = sum8 = sum9 = 0
        for q in qs7:
            sum7 += q.gzB
            if q.xcorwl == True:
                sum8 += q.gzB  # 出席会议的股东所持B股总数
                num8 += 1
            else:
                sum9 += q.gzB
                num9 += 1
        percent7 = sum7 / BShareTotal
        percent8 = sum8 / BShareTotal
        percent9 = sum9 / BShareTotal
        # sum7 = format(sum7, ",")
        # sum8 = format(sum8, ",")
        # sum9 = format(sum9, ",")
        content3 = """        B股股东（代理人）{}人，代表股份{}股，占公司B股股东表决权股份{:.2%}。其中,现场会议股东（代理人）{}人，代表股份{}股，占公司有表决权总股份{:.2%}。网络投标股东（代理人）{}人，代表股份{}股，占公司有表决权总股份{:.2%}。"""\
            .format(num7, sum7, percent7, num8, sum8, percent8, num9, sum9, percent9)
        self.document.add_paragraph(content3)

        self.document.add_paragraph("4. 中小股出席情况：")
        zxg_num = zxg_gb = 0
        for q in qs1:
            s = q.shareholder
            # s = ShareholderInfo.objects.get(gdxm=s)
            # 统计出席会议的中小股信息
            if s.gdtype == "中小股":
                zxg_num += s.rs
                zxg_gb += s.gzB + s.gzA
        content4 = """        中小股东（代理人）（不含公司董事、监事、高管、单独或合计持有公司5%以上股份的股东及与持有公司5%以上股份股东形成一致行动人的股东）共{}人，代表股份{}股，占公司B股股东表决权股份{:.2%}。""".format(zxg_num, zxg_gb, zxg_gb/totalShare)
        self.document.add_paragraph(content4)

        self.document.add_paragraph("二、提案审议和表决情况")
        motion_qs =  m.motionbook_set.all()
        motion_content = []
        # 统计每项议案的表决情况，然后写入
        for i in motion_qs:
            self.document.add_heading(i.name, level=3)
            zancheng = i.zanchengB + i.zanchengA
            zxg_zan = i.zxg_zanchengA + i.zxg_zanchengB
            zxg_fan = i.zxg_fanduiA + i.zxg_fanduiB
            zxg_qi = i.zxg_qiA + i.zxg_qiB
            fandui = i.fanduiA + i.fanduiB
            qiquan = i.qiquanA + i.qiquanB
            percentY = zancheng / cx_gb
            percentN = fandui / cx_gb
            percentQ = qiquan / cx_gb
            percentAY = i.zanchengA/sum4 if sum4 > 0 else 0  # 避免除数为零的情况
            percentAN = i.fanduiA/sum4 if sum4 > 0 else 0
            percentAQ = i.qiquanA/sum4 if sum4 > 0 else 0
            percentBY = i.zanchengB/sum8 if sum8 > 0 else 0
            percentBN = i.fanduiB/sum8 if sum8 > 0 else 0
            percentBQ = i.qiquanB/sum8 if sum8 > 0 else 0
            percent_zxg_z = zxg_zan/zxg_gb if zxg_gb >0 else 0
            percent_zxg_f = zxg_fan/zxg_gb if zxg_gb >0 else 0
            percent_zxg_q = zxg_qi/zxg_gb if zxg_gb >0 else 0

            content = """            (1)总的表决情况：
            同意{}股，占出席会议所有股东所持表决权{:.2%}；反对{}股，占出席会议所有股东所持表决权{:.2%}；弃权{}股，占出席会议所有股东所持表决权{:.2%}。
            (2)A股股东表决情况：
            同意{}股，占出席会议A股东所持表决权{:.2%}；反对{}股，占出席会议A股东所持表决权{:.2%}；弃权{}股，占出席会议A股股东所持表决权{:.2%}。
            (3)B股股东表决情况：
            同意{}股，占出席会议B股东所持表决权{:.2%}；反对{}股，占出席会议A股东所持表决权{:.2%}；弃权{}股，占出席会议B股股东所持表决权{:.2%}。
            (4)中小股东表决情况:
            同意{}股，占出席会议中小股东所持表决权{}%；反对{}股，占出席会议中小股东所持表决权{}%；弃权{}股，占出席会议中小股东所持表决权{}%。
            表决结果："""\
                .format(zancheng, percentY, fandui, percentN, qiquan, percentQ,
               i.zanchengA, percentAY, i.fanduiA, percentAN, i.qiquanA, percentAQ,
                i.zanchengB, percentBY, i.fanduiB, percentBN, i.qiquanB, percentBQ,
                        zxg_zan, percent_zxg_z,zxg_fan, percent_zxg_f,zxg_qi,percent_zxg_q)

            self.document.add_paragraph(content)
            motion_content.append({"name": i.name, "content": content})


        leijimotion_qs = m.accumulatemotion_set.all()
        num = len(leijimotion_qs) # 如果有三个累计议案，则每个股东可投票票数乘以3，总票数也乘以3，在统计的zanchengA和B中已经是乘以3的票数了
        leiji_content = []
        for i in leijimotion_qs:
            self.document.add_heading(i.name, level=3)
            zancheng = i.zanchengB + i.zanchengA
            percentA = i.zanchengA/(sum4*num) if sum4>0 else 0
            percentB = i.zanchengB/(sum4*num) if sum4 > 0 else 0
            content = """            总得票数：{}票，占出席会议所有股东所持表决权{:.2%}; A股股东票数：{}票，占出席会议A股股东所持表决权{:.2%}; B股股东票数：{}票，占出席会议B股股东所持表决权{:.2%}
            表决结果：""".format(zancheng,zancheng/(cx_gb*num),
                            i.zanchengA, percentA,
                            i.zanchengB, percentB)
            self.document.add_paragraph(content)
            leiji_content.append({"name": i.name, "content": content})
        self.document.add_paragraph("唱票统计人签名：")
        self.document.add_paragraph("监票人签名：")

        self.document.save(BASE_DIR + "/files/"+year+name+".doc")
        print(BASE_DIR + "/files/"+year+name+".doc")
        return  {"content1": content1,
                "content2": content2,
                "content3": content3,
                "content4": content4,
                "content5": motion_content,
                "content6": leiji_content
                }
